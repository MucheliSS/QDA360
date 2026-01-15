import re
import pandas as pd
from pathlib import Path
from docx import Document

import logging
logger = logging.getLogger(__name__)

def parse_docx(path: str | Path) -> pd.DataFrame:
    """
    Parse a DOCX transcript into the Qux360 schema.

    Parameters
    ----------
    path : str or Path
        Path to the .docx transcript file.

    Returns
    -------
    pd.DataFrame
        Transcript with columns: timestamp, speaker_id, speaker, statement, codes, themes.
    """
    
    logger.debug(f"Parse DOCX - Path: {path}")
    
    doc = Document(path)
    
    segments = []
    current_speaker = None
    current_timestamp = None
    current_statement = []

    # Helper to process a line of text
    def process_line(line):
        nonlocal current_speaker, current_timestamp, current_statement, segments
        line = line.strip()
        if not line:
            return

        # 1. Try standard pattern: Speaker Name TIMESTAMP
        match = re.match(
            r'([A-Za-z\s!@#$%^&*()_+=\-\{\}\[\]\|;\'\",.<>/?~]+)\s+((?:(\d{1,2}):)?\d{1,2}:\d{2})',
            line
        )

        # 2. Key fallback: "Speaker: Text" pattern (common in cleaned transcripts)
        if not match:
             match_colon = re.match(r'^([A-Za-z0-9\s\-_]+):\s+(.+)', line)
             if match_colon:
                 potential_speaker = match_colon.group(1).strip()
                 if len(potential_speaker) < 50 and len(potential_speaker) > 1:
                     if current_speaker and current_statement:
                        cleaned_statement = " ".join(current_statement).replace("\n", " ")
                        segments.append([current_timestamp, None, current_speaker, cleaned_statement, [], []])
                     
                     current_speaker = potential_speaker
                     current_timestamp = "00:00"
                     current_statement = [match_colon.group(2).strip()]
                     return

        # 3. Implicit Speaker Fallback: Short line that looks like a name (no colon/timestamp)
        # Only strict if we are NOT currently in the middle of a statement? 
        # Actually, often speaker names are just short lines.
        # Heuristic: < 50 chars, no sentence punctuation, distinct from previous text
        if not match and not match_colon:
            is_potential_name = (
                len(line) < 50 
                and not line.endswith(('.', '?', '!', ',', ';'))
                and re.match(r'^[A-Za-z0-9\s\-_]+$', line)
            )
            
            if is_potential_name:
                 # It's likely a speaker name
                 if current_speaker and current_statement:
                    cleaned_statement = " ".join(current_statement).replace("\n", " ")
                    segments.append([current_timestamp, None, current_speaker, cleaned_statement, [], []])
                 
                 current_speaker = line.strip()
                 current_timestamp = "00:00"
                 current_statement = []
                 return

        if match:
            # Save previous segment
            if current_speaker and current_statement:
                cleaned_statement = " ".join(current_statement).replace("\n", " ")
                segments.append([current_timestamp, None, current_speaker, cleaned_statement, [], []])

            # Extract new speaker + timestamp
            current_speaker = match.group(1).strip()
            current_timestamp = match.group(2)
            current_statement = [line.split(current_timestamp, 1)[1].strip()]
        else:
            if current_statement is not None:
                current_statement.append(line)

        if match:
            # Save previous segment
            if current_speaker and current_statement:
                cleaned_statement = " ".join(current_statement).replace("\n", " ")
                segments.append([current_timestamp, None, current_speaker, cleaned_statement, [], []])

            # Extract new speaker + timestamp
            current_speaker = match.group(1).strip()
            current_timestamp = match.group(2)
            current_statement = [line.split(current_timestamp, 1)[1].strip()]
        else:
            if current_statement is not None:
                current_statement.append(line)

    # 1. Process standard body paragraphs
    for para in doc.paragraphs:
        process_line(para.text)

    # 2. Process tables (often used for transcripts)
    for table in doc.tables:
        for row in table.rows:
            # Strategy: Join cell texts? Or treat each cell as a potential paragraph?
            # Common formats:
            # A) Speaker | Text
            # B) Timestamp | Speaker | Text
            # C) Single column with "Speaker: Text"
            
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            
            if not row_text:
                continue

            # Heuristic for multi-column tables
            if len(row_text) >= 2:
                 # Check if first col looks like speaker or timestamp
                 col0 = row_text[0]
                 col1 = row_text[1]
                 
                 # Case: Speaker | Text
                 # If col0 is short (<30 chars) and col1 is longer, assume Speaker | Text
                 if len(col0) < 30 and len(col1) > 0:
                     # Check if it was already processed by previous iteration (unlikely in table)
                     if current_speaker and current_statement:
                        cleaned_statement = " ".join(current_statement).replace("\n", " ")
                        segments.append([current_timestamp, None, current_speaker, cleaned_statement, [], []])
                     
                     current_speaker = col0
                     current_timestamp = "00:00"
                     current_statement = [col1]
                     # If there are more columns, append them to statement
                     if len(row_text) > 2:
                         current_statement.extend(row_text[2:])
                     continue
            
            # Fallback: Process each cell as a paragraph (handles single-col tables or mixed bag)
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_line(para.text)

    # Save last segment
    if current_speaker and current_statement:
        cleaned_statement = " ".join(current_statement).replace("\n", " ")
        segments.append([current_timestamp, None, current_speaker, cleaned_statement, [], []])

    df = pd.DataFrame(
        segments,
        columns=["timestamp", "speaker_id", "speaker", "statement", "codes", "themes"]
    )

    return df
